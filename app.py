from __future__ import annotations

from datetime import date, datetime, timedelta
from io import BytesIO
import re
import json
from pathlib import Path

from flask import Flask, render_template, request, send_file, redirect, url_for, session
import pdfplumber
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

app = Flask(__name__)
app.config["SECRET_KEY"] = "generador-planes-local-2026"
app.config["MAX_CONTENT_LENGTH"] = 12 * 1024 * 1024
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "plantilla" / "FormatoPlanClaseP11-F03.docx"
DAY_NAMES = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]

# Síntesis de los contenidos y actividades del PDF proporcionado.
ACTIVITIES = [
    ("Ruta de aprendizaje front-end", "Apuntes de clase sobre la ruta de aprendizaje de un desarrollador front-end y back-end."),
    ("Editor de texto e IDE", "Video-tutorial de instalación y configuraciones principales de VS Code u otro IDE."),
    ("Atajos y Emmet", "Exposición en equipo sobre atajos de teclado y complementos Emmet."),
    ("HTML y CSS", "Ejercicio individual: maquetar la estructura básica de un sitio web."),
    ("Diseño responsivo", "Crear una página web responsiva para smartphone, tablet y escritorio."),
    ("JSON", "Participar individualmente en el foro sobre JSON y argumentar tres réplicas."),
    ("JavaScript", "Resolver ejercicios de fundamentos, objetos, funciones, clases, asincronía, DOM y APIs."),
    ("Bootstrap: fundamentos", "Elaborar apuntes de clase sobre los fundamentos del framework Bootstrap."),
    ("Bootstrap: sitio web", "Desarrollar un sitio web usando Customize, Layout y Content de Bootstrap."),
    ("Bootstrap: componentes", "Desarrollar un sitio usando contenedores, componentes y formularios de Bootstrap."),
    ("Bootstrap: helpers y utilidades", "Elaborar un mapa mental de Helpers, Utilities y Bootstrap Icons."),
    ("Material UI: fundamentos", "Elaborar un cuadro sinóptico sobre los fundamentos de Material UI."),
    ("Material UI: aplicación", "Crear un sitio web utilizando el framework Material UI."),
    ("Tailwind CSS", "Investigar qué es Tailwind CSS y sus diferencias con Bootstrap."),
    ("Tailwind CSS: práctica", "Desarrollar un sitio web utilizando Tailwind CSS."),
    ("Vue JS: introducción", "Instalar Vue JS y ejecutar la primera aplicación."),
    ("Vue JS: estructura", "Investigar la estructura de un proyecto en Vue JS."),
    ("Vue JS: reactividad", "Implementar reactividad y crear componentes con Vue JS."),
    ("Vue JS: formularios y router", "Crear formulario, agregar validaciones y configurar Vue Router."),
    ("Vue JS: Pinia y Watch", "Practicar el uso de Pinia y Watch para almacenar estados."),
    ("Vue JS: peticiones HTTP", "Realizar peticiones HTTP asíncronas en Vue JS."),
    ("Vue JS: comunicación", "Implementar comunicación entre componentes mediante Props."),
    ("Vue JS: propiedades y directivas", "Usar propiedades computadas, filtros y directivas condicionales e interactivas."),
    ("Proyecto integrador Vue", "Definir y presentar un proyecto integrador tipo CRUD usando Vue."),
]


def scheduled_dates(start: date, end: date, day_indexes: set[int]) -> list[date]:
    current, dates = start, []
    while current <= end:
        if current.weekday() in day_indexes:
            dates.append(current)
        current += timedelta(days=1)
    return dates


def clean_activities(raw: str) -> list[tuple[str, str]]:
    rows = []
    for line in raw.splitlines():
        line = line.strip()
        if line:
            topic, separator, activity = line.partition("|")
            rows.append((topic.strip(), activity.strip() if separator and activity.strip() else "Actividad de aprendizaje por definir."))
    return rows or ACTIVITIES


def extract_learning_activities(upload) -> tuple[list[tuple[str, str]], dict[str, str]]:
    """Convierte la secuencia didáctica a pares Tema (Contenido) / Actividad."""
    with pdfplumber.open(BytesIO(upload.read())) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    text = re.sub(r"[ \t]+", " ", text)
    element_pattern = re.compile(r"Elemento de competencia\s+(\d+)\s*:\s*(.*?)(?=Elemento de competencia\s+\d+\s*:|\Z)", re.I | re.S)
    result, element_names = [], {}
    for element in element_pattern.finditer(text):
        number, block = element.group(1), element.group(2)
        name = re.split(r"Competencias blandas a promover:", block, maxsplit=1, flags=re.I)[0]
        element_names[number] = " ".join(name.split()) or f"Elemento de competencia {number}"
        phase_pattern = re.compile(r"EC\s*" + re.escape(number) + r"\s+Fase\s+[IVX]+\s*:\s*(.*?)(?=\s+EC\s*" + re.escape(number) + r"\s+Fase\s+[IVX]+\s*:|\s+Evaluaci[óo]n formativa:|\Z)", re.I | re.S)
        for phase in phase_pattern.finditer(block):
            phase_block = phase.group(1)
            content_match = re.search(r"Contenido:\s*(.*?)(?=\s+EC\s*" + re.escape(number) + r"\s+F\d+\s+Actividad de aprendizaje|\Z)", phase_block, re.I | re.S)
            content = " ".join(content_match.group(1).split()) if content_match else f"Elemento de competencia {number}"
            activity_pattern = re.compile(r"EC\s*" + re.escape(number) + r"\s+F\d+\s+Actividad de aprendizaje\s+(\d+)\s*:\s*(.*?)(?=\s+Tipo de actividad:|\s+EC\s*" + re.escape(number) + r"\s+F\d+\s+Actividad de aprendizaje|\Z)", re.I | re.S)
            for activity in activity_pattern.finditer(phase_block):
                activity_number = activity.group(1)
                title = " ".join(activity.group(2).split())
                if title.lower().startswith("tipo de actividad:"):
                    title = re.sub(r"^Tipo de actividad:\s*", "", title, flags=re.I)
                    title = re.split(r"\s+Aula\s*\(", title, maxsplit=1, flags=re.I)[0].strip()
                if title:
                    result.append((content, f"EC{number} Actividad de aprendizaje {activity_number}: {title}"))
    return result, element_names


def shade(cell, color: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    props.append(fill)


def set_cell_text(cell, value: str, bold=False, color=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold, run.font.name, run.font.size = bold, "Arial", Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def make_docx(course: str, teacher: str, semester: str, hours: str, sessions: list[dict]) -> BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.left_margin = section.right_margin = Cm(1.6)
    doc.styles["Normal"].font.name, doc.styles["Normal"].font.size = "Arial", Pt(9)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PLAN DE CLASE")
    run.bold, run.font.name, run.font.size = True, "Arial", Pt(15)
    run.font.color.rgb = RGBColor(20, 76, 130)
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(f"Curso: {course or 'Sin especificar'}\n").bold = True
    details.add_run(f"Docente: {teacher or 'Sin especificar'}  |  Periodo: {semester or 'Sin especificar'}  |  Horas por semana: {hours}")
    for session in sessions:
        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        labels = ["Sesión", "Fecha", "Tema", "Actividades de aprendizaje"]
        for cell, label in zip(table.rows[0].cells, labels):
            shade(cell, "164C82")
            set_cell_text(cell, label, True, (255, 255, 255))
        set_cell_text(table.cell(1, 0), str(session["number"]))
        set_cell_text(table.cell(1, 1), session["date"])
        set_cell_text(table.cell(1, 2), session["topic"])
        set_cell_text(table.cell(1, 3), session["activity"])
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Generado con Generador de Planes de Clase").font.size = Pt(8)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def make_template_docx(teacher: str, chief: str, sessions: list[dict], element_names: dict[str, str]) -> BytesIO:
    """Llena una copia DOCX de la plantilla; funciona igual en Windows y Linux."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError("No se encontró la plantilla DOCX del proyecto")
    document = Document(TEMPLATE_PATH)
    rows_by_element = {element: 2 for element in range(1, 5)}
    for session_data in sessions:
        element = session_data["element"]
        if element not in rows_by_element or element > len(document.tables):
            continue
        table = document.tables[element - 1]
        row = rows_by_element[element]
        while row >= len(table.rows):
            table.add_row()
        values = [str(session_data["number"]), session_data["date"], session_data["topic"], session_data["activity"]]
        for column, value in enumerate(values):
            table.cell(row, column).text = value
        rows_by_element[element] += 1
    for element in range(1, min(4, len(document.tables)) + 1):
        element_name = element_names.get(str(element), "")
        document.tables[element - 1].cell(1, 0).text = f"ELEMENTO {element}: {element_name}".strip()
    def replace_in_paragraph(paragraph, old, new):
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, "XXXXXXXXX", teacher or "")
        replace_in_paragraph(paragraph, "XXXXXXXX", chief or "")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, "XXXXXXXXX", teacher or "")
                    replace_in_paragraph(paragraph, "XXXXXXXX", chief or "")
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def extract_schedule(upload) -> dict:
    """Extrae campos administrativos y materias de un horario institucional en PDF."""
    with pdfplumber.open(BytesIO(upload.read())) as pdf:
        page = pdf.pages[0]
        text = page.extract_text() or ""
        words = page.extract_words()
    def capture(pattern: str) -> str:
        match = re.search(pattern, text, re.IGNORECASE)
        return " ".join(match.group(1).split()) if match else ""
    professor = capture(r"NOMBRE DEL PROFESOR:\s*(.*?)\s+NO\.\s*EMPLEADO")
    program = capture(r"PE DE ADSCRIPCI[ÓO]N\s*(.*?)\s+CATEGORIA")
    start = capture(r"FECHA INICIAL:\s*(.*?)\s+FECHA FINAL")
    end = capture(r"FECHA FINAL:\s*(.*?)(?=\s+CLAVE DE MATERIA)")
    # Las líneas de la tabla de materias tienen clave, grupo y horas; se agrupan
    # por coordenada vertical para conservar los nombres que se parten en renglones.
    code_words = [word for word in words if re.fullmatch(r"\d{3}[A-Z]{2}\d{3}", word["text"])]
    courses = []
    for word in code_words:
        y = word["top"]
        same_row = [w for w in words if abs(w["top"] - y) < 2 and w["x0"] > word["x1"]]
        line = " ".join(w["text"] for w in same_row)
        group = re.search(r"\b(\d{3})\b", line)
        if not group:
            continue
        hour_match = re.search(r"\s+(\d+)\s+-\s+(\d+)\s+-\s+(\d+)", line[group.end():])
        name = re.sub(r"\s+\d+\s+-\s+\d+\s+-\s+\d+.*$", "", line[group.end():]).strip()
        weekly_hours = sum(map(int, hour_match.groups())) if hour_match else 0
        courses.append({"code": word["text"], "group": group.group(1), "name": name or "Materia sin identificar", "weekly_hours": weekly_hours, "days": []})
    unique = {(c["code"], c["group"]): c for c in courses}
    courses = list(unique.values())
    day_centers = [137, 219, 301, 383, 465, 547]
    for word in words:
        match = re.fullmatch(r"(\d{3}[A-Z]{2}\d{3})-(\d{3})", word["text"])
        if not match or word["top"] < 180:
            continue
        key = (match.group(1), match.group(2))
        if key in unique:
            nearest_day = min(range(len(day_centers)), key=lambda index: abs(day_centers[index] - word["x0"]))
            if abs(day_centers[nearest_day] - word["x0"]) < 45:
                unique[key]["days"].append(nearest_day)
    for course in courses:
        course["days"] = sorted(set(course["days"]))
    # Nombre de jefatura: se encuentra entre la firma docente y la leyenda de jefatura.
    signature_words = [w for w in words if 700 <= w["top"] <= 718 and w["x0"] >= 350]
    chief_name = " ".join(w["text"] for w in signature_words)
    return {"professor": professor, "program": program, "start_text": start, "end_text": end, "chief": chief_name, "courses": courses}


def spanish_date_to_input(value: str) -> str:
    months = {"enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6, "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12}
    match = re.search(r"(\d{1,2})\s+de\s+([a-záéíóú]+)\s+de\s+(\d{4})", value.lower())
    if not match:
        return ""
    return f"{match.group(3)}-{months.get(match.group(2), 1):02d}-{int(match.group(1)):02d}"


app.add_template_filter(spanish_date_to_input, "date_input")


@app.get("/")
def index():
    return render_template("index.html")


@app.post("/subir-horario")
def upload_schedule():
    upload = request.files.get("schedule")
    if not upload or not upload.filename.lower().endswith(".pdf"):
        return "Sube un archivo PDF de horario.", 400
    try:
        data = extract_schedule(upload)
    except Exception:
        return "No fue posible leer el horario. Verifica que sea un PDF válido.", 400
    session["schedule_data"] = data
    return render_template("secuencia.html", data=data)


@app.post("/subir-secuencia")
def upload_syllabus():
    upload = request.files.get("syllabus")
    if not upload or not upload.filename.lower().endswith(".pdf"):
        return "Sube un PDF de secuencia didáctica.", 400
    try:
        activities, element_names = extract_learning_activities(upload)
    except Exception:
        return "No fue posible leer la secuencia didáctica. Verifica que sea un PDF válido.", 400
    if not activities:
        return "No se detectaron elementos de competencia ni actividades en el PDF.", 400
    data = session.get("schedule_data")
    if not data:
        return redirect(url_for("index"))
    catalog = "\n".join(f"{topic} | {activity}" for topic, activity in activities)
    return render_template("confirmar.html", data=data, catalog=catalog, element_names=element_names)


@app.post("/usar-horario")
def use_schedule():
    course = request.form.get("course", "")
    prefill = {
        "course": course, "teacher": request.form.get("teacher", ""), "program": request.form.get("program", ""),
        "chief": request.form.get("chief", ""), "start_date": request.form.get("start_date", ""),
        "end_date": request.form.get("end_date", ""), "weekly_hours": request.form.get("weekly_hours", ""),
        "selected_days": [int(value) for value in request.form.getlist("days")],
    }
    session["prefill"] = prefill
    return redirect(url_for("index"))


@app.post("/generar")
def generate():
    try:
        start = datetime.strptime(request.form["start_date"], "%Y-%m-%d").date()
        end = datetime.strptime(request.form["end_date"], "%Y-%m-%d").date()
    except (KeyError, ValueError):
        return "Selecciona fechas válidas.", 400
    if end < start:
        return "La fecha de fin debe ser igual o posterior a la fecha de inicio.", 400
    days = {int(day) for day in request.form.getlist("days")}
    if not days:
        return "Selecciona por lo menos un día de clase.", 400
    try:
        hours_by_day = {day: float(request.form.get(f"day_hours_{day}", "0")) for day in days}
    except ValueError:
        return "Indica horas válidas para cada día seleccionado.", 400
    if any(hours <= 0 for hours in hours_by_day.values()):
        return "Indica las horas para cada día seleccionado.", 400
    weekly_hours = sum(hours_by_day.values())
    class_dates = scheduled_dates(start, end, days)
    if not class_dates:
        return "No hay clases en los días seleccionados dentro de ese periodo.", 400
    activities = clean_activities(request.form.get("activities", ""))
    try:
        element_names = json.loads(request.form.get("element_names", "{}"))
    except json.JSONDecodeError:
        element_names = {}
    sessions = [{"number": index, "date": class_date.strftime("%d/%m/%Y"), "topic": activities[(index - 1) % len(activities)][0], "activity": activities[(index - 1) % len(activities)][1]} for index, class_date in enumerate(class_dates, 1)]
    for session_data in sessions:
        match = re.match(r"EC\s*(\d+)", session_data["activity"], re.I)
        session_data["element"] = int(match.group(1)) if match else 1
    try:
        docx = make_template_docx(request.form.get("teacher", ""), request.form.get("chief", ""), sessions, element_names)
    except Exception as exc:
        return f"No fue posible generar el documento desde la plantilla: {exc}", 500
    name = re.sub(r"[^A-Za-z0-9_-]+", "_", request.form.get("course", "plan_clase")).strip("_") or "plan_clase"
    return send_file(docx, as_attachment=True, download_name=f"{name}_plan_de_clase.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    app.run(debug=True, port=5000)
